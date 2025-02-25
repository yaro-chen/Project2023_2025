# -*- coding: utf-8 -*-
"""
Created on Wed Jul 17 18:41:34 2024

@author: yr.chen
"""

from docx import Document
import pandas as pd
import os
import json
import argparse

parser = argparse.ArgumentParser(description="Autofill PFI QC table")
parser.add_argument('-i','--chipID',type=str)
parser.add_argument('-c','--confidence', help='confidence values; seperated by "/" ',type=str)
parser.add_argument('-d','--date', help='date, format: YYYY/MM/DD',type=str)
parser.add_argument('-m','--mainfolder',help= 'Absolute/relative path of the main folder')
parser.add_argument('-DB','--database', help='Name of Database; seperated by "/" ')
args = parser.parse_args() 

filename = "SJ4-TE-029_PFI_QC_V2.0_latest.docx"
print(f"{filename} as the template")
chipID = args.chipID
print(f"Check chip: {chipID}")
date =args.date
print(f"Check date: {date}")
main = os.path.abspath(args.mainfolder)
print(f"Check main folder : {main}")
filepath = os.path.dirname(main)
mainname = os.path.basename(main)

confidence = args.confidence.split('/')
db = args.database.split('/')

dic = dict(zip(db, confidence))
print(f"Database/confidence check: {dic}")

newfile = f'{main}/SJ4-TE-029_PFI_QC_V2.0_{mainname}.docx'

doc = Document(filename)
print(doc.tables)

table=doc.tables[0]
table.cell(1, 2).text = date
table.cell(1, 8).text = chipID

if confidence[0] == confidence[1]:
    table.cell(1, 14).text = confidence[0]
elif confidence[0] != confidence[1]:
    table.cell(1, 14).text = str(dic).replace('{','').replace('}','').replace("'",'')
else :
    table.cell(1,14).text = "N/A"

standard = []
for k,v in dic.items():
    print(f"D6311 result from {k}...")
    command = f'find {main}/{k} -maxdepth 3 -mindepth 2 -name "*D6311*{v}*_sorted.txt" -o -name "*PC*{v}*_sorted.txt" '
    print(f"Command for catching D6311 result : {command}")
    st =  os.popen(command).read().strip()
    if os.path.isfile(st):
        print(f"D6311 bracken.txt for {k} with confidence {v} : {st}")
        standard.append(st)
    else: 
        print(f"{st} does not exist!!") 

# standard check
standard_comp = f"{filepath}/D6311_composition.txt"

with open(standard_comp) as c:
   comp = [row.split('\t')[0] for row in c.readlines()[1:]] 

row = 3
for i in comp:
    reads,fraction = [],[]
    for d in standard:
        df = pd.read_csv(d,sep='\t')
        total = sum(df['new_est_reads'])
        df['precision'] = [ float(row/total) for row in df['new_est_reads'].tolist()]
        if i == "Bacillus subtilis":
            genus = i.split(' ')[0]
            exist = df['name'].str.contains(genus)
            if exist.any():
                check = df[exist]
                reads.append(str(sum(check.loc[:,"new_est_reads"])))
                fraction.append('{:.3g}'.format(sum(check["precision"])*100))
            else:
                check = f"{genus}:not detected"
                reads.append("Undetected")
                fraction.append("Undetected")
        elif i in df.values:
            check = df[df['name'].str.contains(i)]
            reads.append(str(check.loc[:,"new_est_reads"].item()))
            fraction.append('{:.3g}'.format(check.loc[:,"precision"].item()*100))
        else:
            check = f"{i}:not detected"
            reads.append("Undetected")
            fraction.append("Undetected")
    print(check)
    print(reads) 
    print(fraction)

    table.cell(row, 4).text = " / ".join(reads)
    table.cell(row, 10).text = " / ".join(fraction)
    row +=1

# json QC check

QCdir = f"{main}/QC_report"

jlist = []
for j in os.listdir(QCdir):
    if j.endswith(".json"):
        jlist.append(os.path.join(QCdir, j))
print(jlist)

def catching(x,y):
    species = []
    readsN = []
    for k,v in dic.items():
        command_txt = f'find {main}/{k} -maxdepth 3 -mindepth 2 -name "*{k}*{sample}*_{v}_{y}*.txt" '
        print(f"Command for catching {x} results: \n{command_txt} \n")
        target =  os.popen(command_txt).read().strip()
        if os.path.isfile(target):
            print(f"{x} reads for {k} with confidence {v} : \n{st}\n")
            species.append(target)
        else: 
            print(f"{target} does not exist!!")
    for i in species:
        command_row = int(os.popen(f'less {i}|wc -l').read().strip())
        if command_row > 1:
            readsN.append(os.popen(f"tail -n 1 {i}|cut -f6").read().strip())
        elif command_row == 1:
            readsN.append('0')
        else:
            print("Error when checking line number !")
    return readsN

row2 = 14
species = []
for j in jlist:
    with open(j, 'r') as file:
        data = json.load(file)
    sample = os.path.splitext(os.path.basename(j))[0]
    rawreads = '{:3f}M'.format(data['summary']['before_filtering']['total_reads']/1000000)
    cleanreads = '{:3f}M'.format(data['summary']['after_filtering']['total_reads']/1000000)
    cleanQ30 = '{:.3%}'.format(data['summary']['after_filtering']['q30_rate'])

    table.cell(row2, 0).text = sample
    table.cell(row2, 1).text = rawreads
    table.cell(row2, 3).text = cleanreads
    table.cell(row2, 5).text = cleanQ30

# assignable reads check
    bacteriareads = catching('bacteria', 'bacter')
    fungireads = catching('fungi', 'fungi')
    virusreads = catching('virus', 'vir')
 
    table.cell(row2, 9).text = " / ".join(bacteriareads)
    table.cell(row2,11).text = " / ".join(fungireads)
    table.cell(row2,13).text = " / ".join(virusreads)

    row2 +=1

doc.save(newfile)
print(f"save QC table in {newfile}")
